import json
import zipfile

import nbformat
from docx import Document


with open("resultados/metricas.json", encoding="utf-8") as metrics_file:
    metrics = json.load(metrics_file)

recomputed_ratio = 100 * metrics["area_moneda_px2"] / metrics["area_papa_px2"]
assert abs(recomputed_ratio - metrics["area_relativa_pct"]) < 0.01
assert metrics["margen_vs_zona_superior_derecha"] > 0, (
    "La región superior derecha no quedó descartada por el criterio de selección."
)

notebook = nbformat.read("Laboratorio_No_3.ipynb", as_version=4)
nbformat.validate(notebook)
errors = [
    output
    for cell in notebook.cells
    if cell.cell_type == "code"
    for output in cell.outputs
    if output.output_type == "error"
]
assert not errors

document = Document("Reporte_Laboratorio_No_3.docx")
paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
table_text = "\n".join(
    cell.text
    for table in document.tables
    for row in table.rows
    for cell in row.cells
)
assert "Nombre:" in table_text
assert f"{metrics['area_relativa_pct']:.2f}%" in paragraph_text

with zipfile.ZipFile("Reporte_Laboratorio_No_3.docx") as package:
    document_xml = package.read("word/document.xml").decode("utf-8")
    image_count = sum(name.startswith("word/media/") for name in package.namelist())

print(f"Notebook válido: {len(notebook.cells)} celdas, sin errores.")
print(f"Área relativa recomputada: {recomputed_ratio:.4f}%")
print(f"Word válido: {len(document.paragraphs)} párrafos, {image_count} imágenes.")
print(f"Saltos explícitos de página: {document_xml.count('w:type=\"page\"')}")
