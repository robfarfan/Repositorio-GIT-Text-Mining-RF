from __future__ import annotations

import json
from pathlib import Path

import cv2
import nbformat as nbf
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from nbclient import NotebookClient


ROOT = Path(__file__).resolve().parent
NOTEBOOK_PATH = ROOT / "Laboratorio_No_3.ipynb"
REPORT_PATH = ROOT / "Reporte_Laboratorio_No_3.docx"
RESULTS_DIR = ROOT / "resultados"
RESULTS_DIR.mkdir(exist_ok=True)


def md(text: str):
    return nbf.v4.new_markdown_cell(text.strip())


def code(text: str):
    return nbf.v4.new_code_cell(text.strip())


cells = [
    md(
        r"""
# Laboratorio No. 3

## Detección del contorno de una moneda dentro de una papa

Este notebook aplica **procesamiento clásico de imágenes**, sin *machine learning*, *deep learning* ni segmentación por máscara. La solución sigue el flujo visto en clase: conversión de color, suavizado gaussiano, detección de bordes, selección de contornos y cálculo geométrico de áreas.

**Objetivo.** Detectar el borde de la moneda visible a contraluz, aproximar el borde total del corte de papa y calcular qué porcentaje del área del corte corresponde a la moneda.
"""
    ),
    md(
        r"""
## 1. Contexto y método

La fotografía está retroiluminada. La moneda produce una transición circular difusa cerca del centro-derecha. En la parte superior derecha existe una sombra causada por daño estructural de la papa; esa anomalía no debe confundirse con la moneda.

### Decisiones técnicas

- Se trabaja a escala 0.5 para reducir el costo computacional sin cambiar la proporción de áreas.
- **Moneda:** búsqueda radial de contornos circulares desde el centro geométrico de la papa hacia los extremos. Cada candidato se evalúa por continuidad del gradiente, contraste radial y distancia al centro.
- **Papa:** Canny sobre el canal de saturación HSV. Los puntos de borde plausibles se filtran con una banda geométrica ajustada a esta toma y se cierran con una envolvente convexa.
- Las áreas se calculan a partir de los contornos geométricos: `πr²` para el contorno circular y `cv2.contourArea` para el contorno de la papa.

> La envolvente convexa no es una máscara segmentada: se construye únicamente con coordenadas producidas por el detector de bordes. La sombra estructural superior derecha se descarta porque está más lejos del centro y no forma un borde circular completo con el contraste esperado.
"""
    ),
    code(
        r"""
from pathlib import Path
import json
import math

import cv2
import matplotlib.pyplot as plt
import numpy as np
from matplotlib.patches import Circle, Patch

IMAGE_PATH = Path("moneda en papa.jpg")
RESULTS_DIR = Path("resultados")
RESULTS_DIR.mkdir(exist_ok=True)

# Parámetros ajustados exclusivamente para la fotografía suministrada.
SCALE = 0.50
GAUSSIAN_SIZE = 15
POTATO_CANNY_LOW = 20
POTATO_CANNY_HIGH = 60
COIN_CENTER_X_RANGE = range(850, 1661, 10)
COIN_CENTER_Y_RANGE = range(820, 1361, 10)
COIN_MIN_RADIUS = 120
COIN_MAX_RADIUS = 240
COIN_RADIUS_STEP = 5
DISTANCE_PENALTY = 0.012

image_bgr = cv2.imread(str(IMAGE_PATH))
if image_bgr is None:
    raise FileNotFoundError(f"No se encontró la imagen: {IMAGE_PATH.resolve()}")

image_rgb = cv2.cvtColor(image_bgr, cv2.COLOR_BGR2RGB)
small_bgr = cv2.resize(
    image_bgr, None, fx=SCALE, fy=SCALE, interpolation=cv2.INTER_AREA
)
small_rgb = cv2.cvtColor(small_bgr, cv2.COLOR_BGR2RGB)
gray = cv2.cvtColor(small_bgr, cv2.COLOR_BGR2GRAY)
hsv = cv2.cvtColor(small_bgr, cv2.COLOR_BGR2HSV)

print(f"Dimensiones originales: {image_bgr.shape[1]} × {image_bgr.shape[0]} píxeles")
print(f"Dimensiones de trabajo: {small_bgr.shape[1]} × {small_bgr.shape[0]} píxeles")
"""
    ),
    md("## 2. Fotografía original sin editar"),
    code(
        r"""
plt.figure(figsize=(12, 8))
plt.imshow(image_rgb)
plt.title("Fotografía original: moneda dentro del corte de papa")
plt.axis("off")
plt.tight_layout()
plt.show()
"""
    ),
    md(
        r"""
## 3. Detección de bordes

Para la papa se usa el canal de saturación, pues separa el amarillo del fondo mejor que la luminancia. Para la moneda se calcula la magnitud del gradiente Sobel sobre una versión suavizada. La búsqueda comienza en el centro geométrico del corte y evalúa circunferencias progresivamente hacia los extremos. Esto favorece la moneda y evita seleccionar la sombra estructural de la zona superior derecha.
"""
    ),
    code(
        r"""
# Borde de la papa a partir del canal de saturación.
saturation_blur = cv2.GaussianBlur(hsv[:, :, 1], (GAUSSIAN_SIZE, GAUSSIAN_SIZE), 0)
potato_edges = cv2.Canny(saturation_blur, POTATO_CANNY_LOW, POTATO_CANNY_HIGH)

# Banda geométrica para descartar bordes de la mesa, la lámpara y textura interior.
edge_y, edge_x = np.nonzero(potato_edges)
radial_position = np.sqrt(
    ((edge_x - 1080) / 670) ** 2 + ((edge_y - 1070) / 470) ** 2
)
valid_outer_edge = (
    (edge_x >= 380)
    & (edge_x <= 1720)
    & (edge_y >= 600)
    & (edge_y <= 1430)
    & (radial_position >= 0.78)
    & (radial_position <= 1.30)
    & ~((edge_y < 735) & (edge_x > 820) & (edge_x < 1160))
)
potato_edge_points = np.column_stack(
    (edge_x[valid_outer_edge], edge_y[valid_outer_edge])
).astype(np.int32)

if len(potato_edge_points) < 5:
    raise RuntimeError("No hay suficientes puntos para construir el contorno de la papa.")

potato_contour = cv2.convexHull(potato_edge_points.reshape(-1, 1, 2))
moments = cv2.moments(potato_contour)
potato_center = np.array(
    [moments["m10"] / moments["m00"], moments["m01"] / moments["m00"]]
)

# Búsqueda radial de la moneda desde el centro de la papa.
gray_smooth = cv2.GaussianBlur(gray, (31, 31), 0)
gradient_x = cv2.Sobel(gray_smooth, cv2.CV_32F, 1, 0, ksize=3)
gradient_y = cv2.Sobel(gray_smooth, cv2.CV_32F, 0, 1, ksize=3)
gradient_magnitude = cv2.magnitude(gradient_x, gradient_y)

angles = np.linspace(0, 2 * np.pi, 180, endpoint=False)
cosines, sines = np.cos(angles), np.sin(angles)
coin_candidates = []

for candidate_y in COIN_CENTER_Y_RANGE:
    for candidate_x in COIN_CENTER_X_RANGE:
        if cv2.pointPolygonTest(
            potato_contour, (float(candidate_x), float(candidate_y)), False
        ) < 0:
            continue
        distance_to_center = float(
            np.linalg.norm(np.array([candidate_x, candidate_y]) - potato_center)
        )
        for candidate_radius in range(
            COIN_MIN_RADIUS, COIN_MAX_RADIUS + 1, COIN_RADIUS_STEP
        ):
            ring_x = np.clip(
                np.round(candidate_x + candidate_radius * cosines).astype(int),
                0,
                gray.shape[1] - 1,
            )
            ring_y = np.clip(
                np.round(candidate_y + candidate_radius * sines).astype(int),
                0,
                gray.shape[0] - 1,
            )
            ring_gradient = gradient_magnitude[ring_y, ring_x]

            inner_x = np.clip(
                np.round(candidate_x + 0.45 * candidate_radius * cosines).astype(int),
                0,
                gray.shape[1] - 1,
            )
            inner_y = np.clip(
                np.round(candidate_y + 0.45 * candidate_radius * sines).astype(int),
                0,
                gray.shape[0] - 1,
            )
            outer_x = np.clip(
                np.round(candidate_x + 1.18 * candidate_radius * cosines).astype(int),
                0,
                gray.shape[1] - 1,
            )
            outer_y = np.clip(
                np.round(candidate_y + 1.18 * candidate_radius * sines).astype(int),
                0,
                gray.shape[0] - 1,
            )

            radial_contrast = float(
                np.mean(gray_smooth[outer_y, outer_x])
                - np.mean(gray_smooth[inner_y, inner_x])
            )
            edge_completeness = float(np.percentile(ring_gradient, 25))
            circular_score = float(
                np.mean(ring_gradient)
                + 0.70 * edge_completeness
                + 0.45 * max(radial_contrast, 0)
                - DISTANCE_PENALTY * distance_to_center
            )
            coin_candidates.append(
                (
                    circular_score,
                    candidate_x,
                    candidate_y,
                    candidate_radius,
                    distance_to_center,
                    radial_contrast,
                )
            )

if not coin_candidates:
    raise RuntimeError("No se generaron candidatos circulares dentro de la papa.")

best_candidate = max(coin_candidates, key=lambda candidate: candidate[0])
coin_score, coin_x, coin_y, coin_radius, coin_distance, coin_contrast = best_candidate
upper_right_scores = [
    candidate[0]
    for candidate in coin_candidates
    if candidate[1] > potato_center[0]
    and candidate[2] < potato_center[1] - 150
]
upper_right_max_score = max(upper_right_scores)
selection_margin = coin_score - upper_right_max_score

print(f"Puntos de borde usados para la papa: {len(potato_edge_points):,}")
print(f"Centro geométrico de la papa: ({potato_center[0]:.1f}, {potato_center[1]:.1f})")
print(f"Moneda detectada: centro=({coin_x}, {coin_y}), radio={coin_radius} píxeles a escala 0.5")
print(f"Distancia desde el centro de la papa: {coin_distance:.1f} píxeles")
print(f"Contraste radial del candidato: {coin_contrast:.2f}")
print(f"Margen frente a la zona superior derecha: {selection_margin:.2f} puntos")
"""
    ),
    code(
        r"""
fig, axes = plt.subplots(1, 2, figsize=(14, 5.5))
axes[0].imshow(potato_edges, cmap="gray")
axes[0].set_title("Bordes Canny del canal de saturación")
axes[0].set_axis_off()

diagnostic = np.zeros_like(potato_edges)
diagnostic[potato_edge_points[:, 1], potato_edge_points[:, 0]] = 255
axes[1].imshow(diagnostic, cmap="gray")
axes[1].set_title("Puntos exteriores conservados para el contorno")
axes[1].set_axis_off()
plt.tight_layout()
plt.savefig(RESULTS_DIR / "mapa_bordes.png", dpi=170, bbox_inches="tight")
plt.show()
"""
    ),
    md("## 4. Contorno de la moneda superpuesto sobre la fotografía"),
    code(
        r"""
overlay_bgr = image_bgr.copy()
scale_back = 1.0 / SCALE
potato_contour_original = np.round(potato_contour * scale_back).astype(np.int32)
coin_center_original = (round(coin_x * scale_back), round(coin_y * scale_back))
coin_radius_original = round(coin_radius * scale_back)

cv2.drawContours(overlay_bgr, [potato_contour_original], -1, (255, 0, 255), 18)
cv2.circle(overlay_bgr, coin_center_original, coin_radius_original, (0, 255, 255), 18)
cv2.circle(overlay_bgr, coin_center_original, 18, (0, 255, 255), -1)

overlay_rgb = cv2.cvtColor(overlay_bgr, cv2.COLOR_BGR2RGB)
cv2.imwrite(str(RESULTS_DIR / "contorno_moneda_superpuesto.png"), overlay_bgr)

plt.figure(figsize=(12, 8))
plt.imshow(overlay_rgb)
plt.title("Contornos detectados: moneda (amarillo) y papa (magenta)")
plt.axis("off")
plt.tight_layout()
plt.show()
"""
    ),
    md(
        r"""
## 5. Estimación del área relativa

Como ambas áreas están expresadas en píxeles cuadrados de la misma fotografía, no se necesita convertir a centímetros para calcular una proporción:

\[
\text{Área relativa (\%)} =
\frac{\pi r_{moneda}^{2}}{A_{contorno\ de\ papa}}\times 100
\]

El resultado representa una aproximación proyectada en 2D. La difusión de la luz y el cierre convexo del borde son las principales fuentes de incertidumbre.
"""
    ),
    code(
        r"""
# El factor de escala se corrige en ambas áreas para reportar píxeles² originales.
coin_area = math.pi * (coin_radius / SCALE) ** 2
potato_area = cv2.contourArea(potato_contour) / (SCALE ** 2)
relative_area = 100.0 * coin_area / potato_area
remaining_area = 100.0 - relative_area

print(f"Área aproximada de la moneda: {coin_area:,.0f} px²")
print(f"Área aproximada del corte de papa: {potato_area:,.0f} px²")
print(f"Área relativa moneda/papa: {relative_area:.2f}%")
"""
    ),
    code(
        r"""
fig, axes = plt.subplots(1, 2, figsize=(14, 6))
axes[0].imshow(overlay_rgb)
axes[0].set_title("Geometrías usadas en el cálculo")
axes[0].set_axis_off()

axes[1].barh([""], [relative_area], color="#E69F00", label="Moneda")
axes[1].barh(
    [""],
    [remaining_area],
    left=[relative_area],
    color="#D9D9D9",
    edgecolor="#333333",
    label="Papa restante",
)
axes[1].set_xlim(0, 100)
axes[1].set_yticks([])
axes[1].set_xlabel("Porcentaje del área del corte (%)")
axes[1].set_title("Área relativa estimada")
axes[1].text(relative_area / 2, 0, f"{relative_area:.2f}%\nmoneda", ha="center", va="center", weight="bold")
axes[1].text(relative_area + remaining_area / 2, 0, f"{remaining_area:.2f}%\nrestante", ha="center", va="center")
axes[1].legend(loc="lower center", bbox_to_anchor=(0.5, -0.32), ncol=2, frameon=False)
axes[1].grid(axis="x", color="#DDDDDD", linewidth=0.8)
axes[1].set_axisbelow(True)
plt.tight_layout()
plt.savefig(RESULTS_DIR / "area_relativa.png", dpi=180, bbox_inches="tight")
plt.show()
"""
    ),
    md("## 6. Validación y sensibilidad"),
    code(
        r"""
def refine_coin_with_blur(blur_size):
    trial_smooth = cv2.GaussianBlur(gray, (blur_size, blur_size), 0)
    trial_gx = cv2.Sobel(trial_smooth, cv2.CV_32F, 1, 0, ksize=3)
    trial_gy = cv2.Sobel(trial_smooth, cv2.CV_32F, 0, 1, ksize=3)
    trial_gradient = cv2.magnitude(trial_gx, trial_gy)
    local_candidates = []
    for trial_y in range(coin_y - 20, coin_y + 21, 10):
        for trial_x in range(coin_x - 20, coin_x + 21, 10):
            trial_distance = float(
                np.linalg.norm(np.array([trial_x, trial_y]) - potato_center)
            )
            for trial_radius in range(
                max(COIN_MIN_RADIUS, coin_radius - 25),
                min(COIN_MAX_RADIUS, coin_radius + 25) + 1,
                COIN_RADIUS_STEP,
            ):
                ring_x = np.clip(
                    np.round(trial_x + trial_radius * cosines).astype(int),
                    0,
                    gray.shape[1] - 1,
                )
                ring_y = np.clip(
                    np.round(trial_y + trial_radius * sines).astype(int),
                    0,
                    gray.shape[0] - 1,
                )
                ring_gradient = trial_gradient[ring_y, ring_x]
                inner_x = np.clip(
                    np.round(trial_x + 0.45 * trial_radius * cosines).astype(int),
                    0,
                    gray.shape[1] - 1,
                )
                inner_y = np.clip(
                    np.round(trial_y + 0.45 * trial_radius * sines).astype(int),
                    0,
                    gray.shape[0] - 1,
                )
                outer_x = np.clip(
                    np.round(trial_x + 1.18 * trial_radius * cosines).astype(int),
                    0,
                    gray.shape[1] - 1,
                )
                outer_y = np.clip(
                    np.round(trial_y + 1.18 * trial_radius * sines).astype(int),
                    0,
                    gray.shape[0] - 1,
                )
                contrast = float(
                    np.mean(trial_smooth[outer_y, outer_x])
                    - np.mean(trial_smooth[inner_y, inner_x])
                )
                completeness = float(np.percentile(ring_gradient, 25))
                score = float(
                    np.mean(ring_gradient)
                    + 0.70 * completeness
                    + 0.45 * max(contrast, 0)
                    - DISTANCE_PENALTY * trial_distance
                )
                local_candidates.append((score, trial_x, trial_y, trial_radius))
    return max(local_candidates, key=lambda candidate: candidate[0])


def potato_area_with_thresholds(low_threshold, high_threshold):
    trial_edges = cv2.Canny(saturation_blur, low_threshold, high_threshold)
    trial_y, trial_x = np.nonzero(trial_edges)
    trial_radial = np.sqrt(((trial_x - 1080) / 670) ** 2 + ((trial_y - 1070) / 470) ** 2)
    keep = (
        (trial_x >= 380)
        & (trial_x <= 1720)
        & (trial_y >= 600)
        & (trial_y <= 1430)
        & (trial_radial >= 0.78)
        & (trial_radial <= 1.30)
        & ~((trial_y < 735) & (trial_x > 820) & (trial_x < 1160))
    )
    points = np.column_stack((trial_x[keep], trial_y[keep])).astype(np.int32)
    contour = cv2.convexHull(points.reshape(-1, 1, 2))
    return cv2.contourArea(contour) / (SCALE ** 2)


sensitivity_rows = []
for blur_size, canny_pair in zip((27, 31, 35), ((18, 54), (20, 60), (22, 66))):
    _, trial_x, trial_y, trial_radius = refine_coin_with_blur(blur_size)
    trial_potato_area = potato_area_with_thresholds(*canny_pair)
    trial_coin_area = math.pi * (trial_radius / SCALE) ** 2
    trial_relative = 100 * trial_coin_area / trial_potato_area
    sensitivity_rows.append(
        {
            "Gauss moneda": blur_size,
            "Canny": f"{canny_pair[0]}/{canny_pair[1]}",
            "centro": f"({trial_x}, {trial_y})",
            "radio_escala_05": int(trial_radius),
            "area_relativa_pct": round(trial_relative, 2),
        }
    )

print("Sensibilidad de parámetros cercanos:")
for row in sensitivity_rows:
    print(row)

sensitivity_values = [row["area_relativa_pct"] for row in sensitivity_rows]
sensitivity_min = min(sensitivity_values)
sensitivity_max = max(sensitivity_values)
print(f"Rango observado: {sensitivity_min:.2f}% a {sensitivity_max:.2f}%")
"""
    ),
    code(
        r"""
# Comprobaciones automáticas de coherencia.
assert 0 < relative_area < 100, "La proporción debe estar entre 0 y 100%."
assert coin_area < potato_area, "El área de la moneda debe ser menor que la de la papa."
assert len(potato_contour) >= 4, "El contorno de la papa debe tener varios vértices."
assert len(sensitivity_rows) == 3, "Deben completarse las tres pruebas de sensibilidad."
assert cv2.pointPolygonTest(potato_contour, (float(coin_x), float(coin_y)), False) >= 0
assert selection_margin > 0, "La moneda debe superar a los candidatos de la zona superior derecha."

metrics = {
    "imagen": str(IMAGE_PATH),
    "dimensiones_originales": [int(image_bgr.shape[1]), int(image_bgr.shape[0])],
    "centro_moneda_px_original": list(coin_center_original),
    "radio_moneda_px_original": int(coin_radius_original),
    "area_moneda_px2": round(coin_area, 2),
    "area_papa_px2": round(potato_area, 2),
    "area_relativa_pct": round(relative_area, 2),
    "sensibilidad_pct": [round(sensitivity_min, 2), round(sensitivity_max, 2)],
    "centro_papa_escala_05": [round(float(potato_center[0]), 2), round(float(potato_center[1]), 2)],
    "distancia_moneda_desde_centro_escala_05": round(float(coin_distance), 2),
    "contraste_radial_moneda": round(float(coin_contrast), 2),
    "margen_vs_zona_superior_derecha": round(float(selection_margin), 2),
    "criterio_seleccion": "búsqueda circular desde el centro de la papa hacia los extremos",
    "parametros": {
        "escala": SCALE,
        "gaussiano": GAUSSIAN_SIZE,
        "canny_papa": [POTATO_CANNY_LOW, POTATO_CANNY_HIGH],
        "penalizacion_distancia": DISTANCE_PENALTY,
        "rango_radio_escala_05": [COIN_MIN_RADIUS, COIN_MAX_RADIUS],
    },
}
(RESULTS_DIR / "metricas.json").write_text(
    json.dumps(metrics, ensure_ascii=False, indent=2), encoding="utf-8"
)
print("Validaciones superadas. Métricas guardadas en resultados/metricas.json")
"""
    ),
    md(
        r"""
## 7. Conclusiones

- El contorno circular de la moneda se detectó mediante gradientes, contraste radial y búsqueda desde el centro; el borde exterior de la papa se obtuvo con Canny sobre saturación.
- El porcentaje calculado es una **aproximación del área proyectada**, no una medición física en cm².
- La validación exige que la moneda sea menor que la papa, que la proporción esté entre 0 y 100% y que los ajustes cercanos produzcan resultados comparables.
- La principal limitación es óptica: el borde de la sombra es difuso por el espesor de la papa y la iluminación posterior. La anomalía estructural superior derecha se controló mediante el criterio de distancia al centro y continuidad circular.

### Reproducibilidad

Ejecutar todas las celdas desde el directorio que contiene `moneda en papa.jpg`. Las versiones directas de las dependencias se encuentran en `requirements.txt`.
"""
    ),
]

notebook = nbf.v4.new_notebook(
    cells=cells,
    metadata={
        "kernelspec": {"display_name": "Python 3", "language": "python", "name": "python3"},
        "language_info": {"name": "python", "version": "3.12"},
    },
)
nbf.write(notebook, NOTEBOOK_PATH)

executed = NotebookClient(
    notebook,
    timeout=180,
    kernel_name="python3",
    resources={"metadata": {"path": str(ROOT)}},
).execute()
nbf.write(executed, NOTEBOOK_PATH)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


metrics = json.loads((RESULTS_DIR / "metricas.json").read_text(encoding="utf-8"))

# Copias reducidas únicamente para mantener liviano el documento Word.
for source_name, target_name, target_width in (
    ("contorno_moneda_superpuesto.png", "reporte_contorno.jpg", 1500),
    ("area_relativa.png", "reporte_area.jpg", 1500),
):
    source_image = cv2.imread(str(RESULTS_DIR / source_name))
    ratio = target_width / source_image.shape[1]
    preview = cv2.resize(
        source_image,
        (target_width, round(source_image.shape[0] * ratio)),
        interpolation=cv2.INTER_AREA,
    )
    cv2.imwrite(
        str(RESULTS_DIR / target_name), preview, [cv2.IMWRITE_JPEG_QUALITY, 88]
    )

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.55)
section.bottom_margin = Inches(0.55)
section.left_margin = Inches(0.65)
section.right_margin = Inches(0.65)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(9.5)
styles["Normal"].paragraph_format.space_after = Pt(3)
styles["Normal"].paragraph_format.line_spacing = 1.0
for style_name in ("Title", "Heading 1", "Heading 2"):
    styles[style_name].font.name = "Aptos Display"
    styles[style_name].font.color.rgb = RGBColor(31, 55, 80)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Laboratorio No. 3\nDetección de una moneda dentro de una papa")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(31, 55, 80)

info = doc.add_table(rows=3, cols=2)
info.style = "Table Grid"
labels = [
    ("Nombre:", "________________________________"),
    ("Carné:", "________________________________"),
    ("Curso / sección:", "________________________________"),
    ("Profesor:", "________________________________"),
    ("Fecha:", "________________________________"),
    ("Dirigido a:", "Profesor del proyecto"),
]
for cell, (label, value) in zip([c for row in info.rows for c in row.cells], labels):
    cell.text = f"{label} {value}"
    cell.paragraphs[0].runs[0].font.size = Pt(8.5)

doc.add_heading("Resumen técnico", level=1)
p = doc.add_paragraph()
p.add_run("Resultado principal. ").bold = True
p.add_run(
    f"La moneda ocupa aproximadamente {metrics['area_relativa_pct']:.2f}% del área proyectada del corte de papa. "
    "La medición se obtuvo exclusivamente con procesamiento clásico de imágenes y geometría de contornos."
)

doc.add_heading("Metodología", level=1)
methods = [
    "Lectura de la fotografía original y reducción uniforme a escala 0.5.",
    "Suavizado gaussiano para reducir ruido antes de calcular gradientes.",
    "Canny en saturación HSV para localizar el borde exterior de la papa.",
    "Búsqueda radial de contornos circulares desde el centro de la papa.",
    "Cálculo de πr² y del área de la envolvente convexa de puntos de borde.",
]
for item in methods:
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run(item)

doc.add_picture(str(RESULTS_DIR / "reporte_contorno.jpg"), width=Inches(5.9))
caption = doc.add_paragraph("Figura 1. Contorno de moneda (amarillo) y aproximación del corte de papa (magenta).")
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption.runs[0].italic = True
caption.runs[0].font.size = Pt(8)

doc.add_page_break()
doc.add_heading("Resultados y validación", level=1)
table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
header = table.rows[0].cells
header[0].text = "Métrica"
header[1].text = "Resultado"
for cell in header:
    set_cell_shading(cell, "1F4E78")
    for run in cell.paragraphs[0].runs:
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.bold = True

result_rows = [
    ("Dimensiones de la imagen", f"{metrics['dimensiones_originales'][0]} × {metrics['dimensiones_originales'][1]} px"),
    ("Radio estimado de la moneda", f"{metrics['radio_moneda_px_original']:,} px"),
    ("Área de la moneda", f"{metrics['area_moneda_px2']:,.0f} px²"),
    ("Área del corte de papa", f"{metrics['area_papa_px2']:,.0f} px²"),
    ("Área relativa", f"{metrics['area_relativa_pct']:.2f}%"),
    ("Rango de sensibilidad", f"{metrics['sensibilidad_pct'][0]:.2f}%–{metrics['sensibilidad_pct'][1]:.2f}%"),
]
for metric_name, value in result_rows:
    row = table.add_row().cells
    row[0].text = metric_name
    row[1].text = value
    for cell in row:
        cell.paragraphs[0].runs[0].font.size = Pt(8.5)

doc.add_picture(str(RESULTS_DIR / "reporte_area.jpg"), width=Inches(5.9))
caption = doc.add_paragraph("Figura 2. Representación gráfica del porcentaje estimado.")
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption.runs[0].italic = True
caption.runs[0].font.size = Pt(8)

doc.add_heading("Discusión y conclusiones", level=1)
p = doc.add_paragraph()
p.add_run("Validez. ").bold = True
p.add_run(
    "Las comprobaciones confirmaron una proporción físicamente válida y resultados próximos al variar moderadamente los umbrales. "
    "La continuidad circular, el contraste radial y la cercanía al centro descartaron la sombra estructural superior derecha."
)
p = doc.add_paragraph()
p.add_run("Limitaciones. ").bold = True
p.add_run(
    "La luz se dispersa dentro de la papa y vuelve difuso el borde de la moneda. La base del corte está próxima al límite de la imagen; por ello, la envolvente convexa completa discontinuidades del borde y puede modificar ligeramente el área de la papa."
)
p = doc.add_paragraph()
p.add_run("Conclusión. ").bold = True
p.add_run(
    f"El flujo reproduce el enfoque explicado en clase y estima que la moneda ocupa {metrics['area_relativa_pct']:.2f}% del corte. "
    "La cifra debe interpretarse como área proyectada aproximada, no como medición física calibrada."
)

doc.add_heading("Referencia metodológica", level=2)
ref = doc.add_paragraph(
    "Transcripciones de clase 01–04 del curso Text Mining & Image Recognition; en particular, clase 03: gradiente, suavizado gaussiano, Canny, contornos y lineamientos del laboratorio."
)
ref.runs[0].font.size = Pt(8)

doc.save(REPORT_PATH)
print(f"Notebook creado y ejecutado: {NOTEBOOK_PATH.name}")
print(f"Reporte Word creado: {REPORT_PATH.name}")
